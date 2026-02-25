"""
방식 A: python-docx를 이용한 표 생성
- 한글(HWP) 미설치 환경에서도 동작
- 생성된 .docx를 한글에서 열어 .hwp로 저장하면 됨

설치: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────
# 헬퍼: 셀 배경색 설정
# ──────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ──────────────────────────────────────────
# 헬퍼: 셀 테두리 설정
# ──────────────────────────────────────────
def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), val.get('val', 'single'))
            b.set(qn('w:sz'), str(val.get('sz', 6)))
            b.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(b)
    tcPr.append(tcBorders)

# ──────────────────────────────────────────
# 메인: 표 생성
# ──────────────────────────────────────────
def create_feasibility_table(output_path: str = "feasibility_v7_table.docx"):
    doc = Document()

    # 페이지 여백 줄이기 (A4 기준)
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── 섹션 제목 ──────────────────────────────
    title = doc.add_paragraph("2-2. 고객 요구사항 대응 방안")
    title.style.font.size = Pt(13)
    title.style.font.bold = True

    subtitle = doc.add_paragraph("■ 압도적 가성비와 무결점 보안성으로 시장(중소 제약사) 수요 정조준")
    subtitle.style.font.size = Pt(11)
    subtitle.style.font.bold = True

    # ── 표 데이터 정의 ──────────────────────────
    TABLE_DATA = [
        # (헤더 여부, 요구사항, 해결방안)
        (True,  "요구사항 (Pain)",    "자사 솔루션 해결 방안"),
        (False, "1. 비용 장벽\n(가성비)",
                "솔루션 초기 구축비용을 기존 해외 엔진 대비 1/10 수준으로 획기적 절감. "
                "고도화된 연산은 사내 PC 분산 그리딩으로 전기세 수준으로 타격 흡수. "
                "추가로 글로벌 검증 모델 업데이트 보증을 통해 글로벌 최전선 유지."),
        (False, "2. 태생적 불안감\n(보안성)",
                "AWS, GCP 등 외부 퍼블릭 클라우드로 1바이트의 정보도 나가지 않으며, "
                "회사 내부 IP 주소로만 접근되는 100% 닫힌 폐쇄 생태계를 지원합니다."),
    ]

    # ── 표 생성 (3행 2열) ──────────────────────
    table = doc.add_table(rows=len(TABLE_DATA), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 열 너비 설정 (30% : 70%)
    total_width = Cm(16)
    table.columns[0].width = int(total_width * 0.30)
    table.columns[1].width = int(total_width * 0.70)

    HEADER_BG = "333333"   # 진한 회색 (헤더)
    ODD_BG    = "F7F7F7"   # 연한 회색 (홀수 행)
    EVEN_BG   = "FFFFFF"   # 흰색 (짝수 행)
    BORDER    = {"val": "single", "sz": 6, "color": "333333"}

    for row_idx, (is_header, col0_text, col1_text) in enumerate(TABLE_DATA):
        row = table.rows[row_idx]
        row.height = Cm(1.2) if is_header else Cm(2.0)

        for col_idx, (cell, text) in enumerate([(row.cells[0], col0_text), (row.cells[1], col1_text)]):
            # 배경색
            if is_header:
                set_cell_bg(cell, HEADER_BG)
            else:
                set_cell_bg(cell, ODD_BG if row_idx % 2 == 1 else EVEN_BG)

            # 테두리
            set_cell_border(cell, top=BORDER, bottom=BORDER, left=BORDER, right=BORDER)

            # 수직 정렬
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 텍스트
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if (col_idx == 0) else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(text)
            run.font.size = Pt(10)
            run.font.bold = is_header or (col_idx == 0)
            if is_header:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # 흰색 글자

    # ── 저장 ────────────────────────────────────
    doc.save(output_path)
    print(f"✅ 저장 완료: {output_path}")
    print("   → 한글(HWP)에서 파일을 열어 '다른 이름으로 저장' → .hwp 형식으로 저장하세요.")


if __name__ == "__main__":
    create_feasibility_table("feasibility_v7_table.docx")
